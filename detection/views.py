from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, authenticate
from django.contrib.auth.forms import UserCreationForm
from django.contrib import messages
from .models import Cattle, TestResult
from django.contrib.auth.models import User
from .utils import run_prediction
from django.core.paginator import Paginator
from django.db.models import Count, Q
from datetime import datetime
from django.core.files.storage import default_storage
from django.db.models.functions import TruncDate
from django.http import JsonResponse, HttpResponse
import torch
from torchvision import transforms
from PIL import Image
import os
import logging
from django.template.loader import render_to_string
from django.conf import settings
import io

# PDF and Word document generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Import the fallback explainability function
from .explainability_fallback import generate_explainability_images_fallback

# Set up logging
logger = logging.getLogger(__name__)

# Load the ONNX model with better error handling
try:
    import onnxruntime as ort
    model_path = os.path.join(os.path.dirname(__file__), '..', 'cattle_disease_model.onnx')
    if os.path.exists(model_path):
        model = ort.InferenceSession(model_path)
        logger.info("ONNX model loaded successfully")
    else:
        logger.warning(f"ONNX model not found at {model_path}")
        model = None
except Exception as e:
    logger.error(f"Error loading ONNX model: {e}")
    model = None

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

def signup(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Account created successfully! Please log in.')
            return redirect('login')
        else:
            messages.error(request, 'There was an error with your submission.')
    else:
        form = UserCreationForm()
    return render(request, 'registration/signup.html', {'form': form})

@login_required
def dashboard(request):
    return render(request, 'detection/dashboard.html')

def dashboard_data(request):
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Authentication required'}, status=401)

    user = request.user
    
    # Get user's cattle and test results
    user_cattle = Cattle.objects.filter(owner=user)
    user_test_results = TestResult.objects.filter(cattle__owner=user)
    
    cattle_count = user_cattle.count()
    total_tests = user_test_results.count()
    healthy_count = user_test_results.filter(is_healthy=True).count()
    infected_count = user_test_results.filter(is_infected=True).count()
    foot_infections = user_test_results.filter(foot_infection=True).count()
    mouth_infections = user_test_results.filter(mouth_infection=True).count()

    # Get trend data
    trend_data = (
        user_test_results
        .values('tested_at__date')
        .annotate(
            healthy_count=Count('id', filter=Q(is_healthy=True)),
            foot_count=Count('id', filter=Q(foot_infection=True)),
            mouth_count=Count('id', filter=Q(mouth_infection=True))
        )
        .order_by('tested_at__date')
    )

    trend_dates = [entry['tested_at__date'].strftime('%Y-%m-%d') for entry in trend_data]
    trend_healthy = [entry['healthy_count'] for entry in trend_data]
    trend_foot = [entry['foot_count'] for entry in trend_data]
    trend_mouth = [entry['mouth_count'] for entry in trend_data]

    # Provide default data if no trends exist
    if not trend_dates:
        from datetime import date
        trend_dates = [date.today().strftime('%Y-%m-%d')]
        trend_healthy = [healthy_count]
        trend_foot = [foot_infections]
        trend_mouth = [mouth_infections]

    data = {
        'cattle_count': cattle_count,
        'total_tests': total_tests,
        'healthy_count': healthy_count,
        'infected_count': infected_count,
        'foot_infections': foot_infections,
        'mouth_infections': mouth_infections,
        'trend_dates': trend_dates,
        'trend_healthy': trend_healthy,
        'trend_foot': trend_foot,
        'trend_mouth': trend_mouth,
    }
    return JsonResponse(data)

@login_required
def upload_image(request):
    if request.method == 'POST':
        image = request.FILES.get('image')
        cattle_id = request.POST.get('cattle_id')
        
        if not image or not cattle_id:
            messages.error(request, "Please select both an image and cattle.")
            return redirect('upload_image')
            
        try:
            # Verify cattle ownership
            cattle = Cattle.objects.get(id=cattle_id, owner=request.user)
            
            # Create TestResult instance
            test_result = TestResult.objects.create(
                cattle=cattle,
                image=image,
                user=request.user
            )
            
            # Process the image
            success = process_image_and_predict(test_result)
            
            if success:
                messages.success(request, "Image processed successfully!")
                return redirect('results', test_id=test_result.id)
            else:
                messages.error(request, "Error processing image. Please try again.")
                test_result.delete()
                
        except Cattle.DoesNotExist:
            messages.error(request, "Invalid cattle selection.")
        except Exception as e:
            logger.error(f"Error in upload_image: {e}")
            messages.error(request, f"Error: {str(e)}")
    
    cattle = Cattle.objects.filter(owner=request.user)
    return render(request, 'detection/upload.html', {'cattle': cattle})

def process_image_and_predict(test_result):
    """
    Process the uploaded image and run prediction with fallback explainability
    """
    try:
        # Get the image path
        image_path = test_result.image.path
        
        # Run basic prediction using utils
        prediction = run_prediction(image_path)
        
        # Update TestResult with prediction results
        test_result.foot_infection = prediction.get('infected_foot', False)
        test_result.mouth_infection = prediction.get('infected_mouth', False)
        test_result.is_infected = prediction.get('is_infected', False)
        test_result.is_healthy = prediction.get('is_healthy', True)
        
        # Generate explainability images using fallback method
        try:
            explainability_success = generate_explainability_images_fallback(image_path, test_result)
            if explainability_success:
                logger.info(f"Generated explainability images for test result {test_result.id}")
            else:
                logger.warning(f"Failed to generate some explainability images for test result {test_result.id}")
        except Exception as e:
            logger.error(f"Error generating explainability images: {e}")
        
        test_result.save()
        return True
        
    except Exception as e:
        logger.error(f"Error in process_image_and_predict: {e}")
        return False

@login_required
def results(request, test_id):
    try:
        test_result = TestResult.objects.get(id=test_id, user=request.user)
        context = {
            'test_result': test_result,
            'image_url': test_result.image.url,
            'integrated_gradients_url': test_result.integrated_gradients_image.url if test_result.integrated_gradients_image else None,
            'gradcam_url': test_result.gradcam_image.url if test_result.gradcam_image else None,
            'occlusion_url': test_result.occlusion_image.url if test_result.occlusion_image else None,
            'lime_url': test_result.lime_image.url if test_result.lime_image else None,
        }
        return render(request, 'detection/results.html', context)
    except TestResult.DoesNotExist:
        messages.error(request, "Test result not found.")
        return redirect('dashboard')

@login_required
def history(request):
    test_results = TestResult.objects.filter(user=request.user).order_by('-tested_at')
    paginator = Paginator(test_results, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, 'detection/history.html', {'page_obj': page_obj})

@login_required
def register_cattle(request):
    if request.method == 'POST':
        tag_number = request.POST.get('tag_number')
        breed = request.POST.get('breed')
        age = request.POST.get('age')
        
        if not all([tag_number, breed, age]):
            messages.error(request, "All fields are required.")
            return render(request, 'detection/register_cattle.html')
            
        try:
            # Check if tag number already exists for this user
            if Cattle.objects.filter(tag_number=tag_number, owner=request.user).exists():
                messages.error(request, "A cattle with this tag number already exists.")
                return render(request, 'detection/register_cattle.html')
                
            Cattle.objects.create(
                tag_number=tag_number,
                breed=breed,
                age=int(age),
                owner=request.user
            )
            messages.success(request, "Cattle registered successfully.")
            return redirect('dashboard')
        except ValueError:
            messages.error(request, "Age must be a valid number.")
        except Exception as e:
            logger.error(f"Error registering cattle: {e}")
            messages.error(request, f"Error: {str(e)}")
            
    return render(request, 'detection/register_cattle.html')

def register_user(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            messages.success(request, "Registration successful.")
            return redirect('dashboard')
        else:
            messages.error(request, "Registration failed. Please check the form.")
    else:
        form = UserCreationForm()
    return render(request, 'detection/register.html', {'form': form})

@login_required
def admin_dashboard(request):
    if not request.user.is_superuser:
        messages.error(request, "You do not have permission to access this page.")
        return redirect('dashboard')
        
    users = User.objects.all().order_by('-date_joined')
    test_results = TestResult.objects.all().order_by('-tested_at')[:50]  # Limit for performance
    
    # Get some statistics
    total_users = users.count()
    total_tests = TestResult.objects.count()
    total_cattle = Cattle.objects.count()
    
    context = {
        'users': users,
        'test_results': test_results,
        'total_users': total_users,
        'total_tests': total_tests,
        'total_cattle': total_cattle,
    }
    return render(request, 'detection/admin_dashboard.html', context)

# NEW DOWNLOAD FUNCTIONALITY BELOW

@login_required
def download_user_guide(request):
    """
    View to handle user guide downloads
    """
    format_type = request.GET.get('format', 'pdf')
    
    if format_type == 'pdf':
        return generate_pdf_guide(request)
    elif format_type == 'docx':
        return generate_docx_guide(request)
    else:
        messages.error(request, "Invalid download format requested.")
        return redirect('dashboard')

def generate_pdf_guide(request):
    """
    Generate PDF version of the user guide
    """
    try:
        # Create a BytesIO buffer to receive PDF data
        buffer = io.BytesIO()
        
        # Create the PDF object using ReportLab
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for the 'Flowable' objects
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#1e40af')
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.HexColor('#2563eb')
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=14,
            spaceAfter=10,
            textColor=colors.HexColor('#3b82f6')
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_JUSTIFY
        )
        
        # Title page
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph("🐄 Cattle Disease Detection System", title_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("Comprehensive User Guide & System Documentation", styles['Heading2']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("AI-Powered Disease Detection for Modern Cattle Management", styles['Normal']))
        story.append(Spacer(1, 1*inch))
        story.append(Paragraph(f"Generated for: {request.user.username}", styles['Normal']))
        story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(PageBreak())
        
        # Table of Contents
        story.append(Paragraph("📋 Table of Contents", heading_style))
        toc_data = [
            ["1.", "System Overview", "3"],
            ["2.", "Key Features", "4"],
            ["3.", "Getting Started", "5"],
            ["4.", "User Interaction Flow", "6"],
            ["5.", "User Interface Guide", "8"],
            ["6.", "Device Compatibility", "9"],
            ["7.", "Security & Privacy", "10"],
            ["8.", "Benefits for Users", "11"],
            ["9.", "Troubleshooting", "12"],
            ["10.", "Support & Contact", "13"],
        ]
        
        toc_table = Table(toc_data, colWidths=[0.5*inch, 4*inch, 1*inch])
        toc_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(toc_table)
        story.append(PageBreak())
        
        # System Overview
        story.append(Paragraph("🎯 System Overview", heading_style))
        story.append(Paragraph(
            "The Cattle Disease Detection System is an advanced AI-powered web application designed to help farmers, "
            "veterinarians, and livestock managers detect diseases in cattle through intelligent image analysis. "
            "The system leverages cutting-edge machine learning algorithms to analyze uploaded images and identify "
            "potential foot and mouth infections in cattle with high accuracy.",
            normal_style
        ))
        story.append(Spacer(1, 12))
        
        # Technology Foundation
        story.append(Paragraph("🔬 Technology Foundation", subheading_style))
        tech_data = [
            ["AI Engine", "ONNX-based machine learning models"],
            ["Image Processing", "Computer vision algorithms"],
            ["Web Platform", "Django-based responsive web application"],
            ["Database", "PostgreSQL for secure data storage"],
            ["Visualization", "Multiple AI explanation techniques"],
        ]
        
        tech_table = Table(tech_data, colWidths=[2*inch, 4*inch])
        tech_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#374151')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#f9fafb')]),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d1d5db')),
        ]))
        story.append(tech_table)
        story.append(PageBreak())
        
        # Key Features
        story.append(Paragraph("⭐ Key Features", heading_style))
        
        features = [
            ("🤖 AI-Powered Detection", "Advanced computer vision models analyze cattle images with professional-grade accuracy"),
            ("⚡ Real-time Analysis", "Get instant results with detailed explanations and visual indicators"),
            ("📊 Comprehensive Analytics", "Dashboard with health statistics, trends, and interactive charts"),
            ("🔍 Visual Explanations", "AI-generated heatmaps and analysis showing areas of concern"),
            ("📱 Mobile Responsive", "Access from any device - desktop, tablet, or smartphone"),
            ("🔒 Secure & Private", "Enterprise-grade security with encrypted data storage"),
        ]
        
        for title, description in features:
            story.append(Paragraph(f"<b>{title}</b>", normal_style))
            story.append(Paragraph(description, normal_style))
            story.append(Spacer(1, 6))
        
        story.append(PageBreak())
        
        # Getting Started
        story.append(Paragraph("🚀 Getting Started", heading_style))
        
        story.append(Paragraph("📝 Account Registration", subheading_style))
        registration_steps = [
            "Navigate to the system URL and click 'Create an account' on the login page",
            "Enter a unique username and create a strong password (minimum 8 characters)",
            "Confirm your password and review password strength indicators",
            "Your account is immediately active after successful registration"
        ]
        
        for i, step in enumerate(registration_steps, 1):
            story.append(Paragraph(f"{i}. {step}", normal_style))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph("🔐 System Login", subheading_style))
        story.append(Paragraph(
            "Access the system using your registered credentials. The login page features password visibility toggle "
            "and secure session management. Sessions automatically expire for security purposes.",
            normal_style
        ))
        
        story.append(PageBreak())
        
        # User Workflow
        story.append(Paragraph("👤 User Interaction Flow", heading_style))
        
        story.append(Paragraph("🏠 Dashboard Overview", subheading_style))
        story.append(Paragraph(
            "Upon login, users are greeted with a comprehensive dashboard that provides real-time health metrics, "
            "interactive charts, quick actions, and recent activity updates. The dashboard serves as the central "
            "hub for all cattle health monitoring activities.",
            normal_style
        ))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph("🐄 Cattle Registration Process", subheading_style))
        cattle_steps = [
            "Access 'Register Cattle' from the sidebar menu",
            "Enter cattle information: Tag Number, Breed, and Age",
            "System validates unique tag numbers and required fields",
            "Successful registration displays confirmation and redirects to dashboard"
        ]
        
        for i, step in enumerate(cattle_steps, 1):
            story.append(Paragraph(f"{i}. {step}", normal_style))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph("🔬 Disease Detection Process", subheading_style))
        detection_steps = [
            "Upload image using drag & drop or file browser",
            "Select cattle from registered animals dropdown",
            "AI processes image and generates analysis",
            "View comprehensive results with visual explanations",
            "Take appropriate action based on findings"
        ]
        
        for i, step in enumerate(detection_steps, 1):
            story.append(Paragraph(f"{i}. {step}", normal_style))
        
        story.append(PageBreak())
        
        # Security & Privacy
        story.append(Paragraph("🔒 Security & Privacy", heading_style))
        
        story.append(Paragraph("🛡️ Data Protection Measures", subheading_style))
        security_features = [
            "Encrypted Authentication: Secure login with password hashing",
            "Data Isolation: User data completely separated and protected",
            "Secure Image Storage: Uploaded images encrypted at rest",
            "HTTPS Encryption: All data transmission encrypted",
            "Session Management: Automatic logout for inactive sessions"
        ]
        
        for feature in security_features:
            story.append(Paragraph(f"• {feature}", normal_style))
        
        story.append(PageBreak())
        
        # Benefits
        story.append(Paragraph("🎯 Benefits for Users", heading_style))
        
        story.append(Paragraph("💰 Economic Benefits", subheading_style))
        economic_benefits = [
            "Cost Savings: Reduce veterinary costs through early detection",
            "Time Efficiency: Quick analysis compared to manual methods",
            "Productivity Gains: Maintain healthier herds",
            "Targeted Treatment: Focus resources where needed"
        ]
        
        for benefit in economic_benefits:
            story.append(Paragraph(f"• {benefit}", normal_style))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph("🔬 Technical Benefits", subheading_style))
        technical_benefits = [
            "Early Detection: Identify diseases before visible symptoms",
            "Accuracy: AI analysis reduces human error",
            "Consistency: Standardized analysis across all cattle",
            "Documentation: Automatic digital health records"
        ]
        
        for benefit in technical_benefits:
            story.append(Paragraph(f"• {benefit}", normal_style))
        
        story.append(PageBreak())
        
        # Support
        story.append(Paragraph("📞 Support & Contact", heading_style))
        story.append(Paragraph(
            "For technical support, general inquiries, or training requests, please contact our support team. "
            "The system includes comprehensive help resources, video tutorials, and best practices guides.",
            normal_style
        ))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph("⚠️ Important Disclaimer", subheading_style))
        story.append(Paragraph(
            "This system is designed to assist in cattle health monitoring but should not replace professional "
            "veterinary diagnosis. Always consult with a qualified veterinarian for serious health concerns.",
            normal_style
        ))
        
        # Footer
        story.append(Spacer(1, 1*inch))
        story.append(Paragraph("Cattle Disease Detection System - Version 2.0", styles['Normal']))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Paragraph("© 2024 Cattle Health Technologies. All rights reserved.", styles['Normal']))
        
        # Build PDF
        doc.build(story)
        
        # Get the value of the BytesIO buffer and write it to the response
        pdf = buffer.getvalue()
        buffer.close()
        
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Cattle_Disease_Detection_User_Guide_{datetime.now().strftime("%Y%m%d")}.pdf"'
        response.write(pdf)
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        messages.error(request, "Error generating PDF. Please try again.")
        return redirect('dashboard')

def generate_docx_guide(request):
    """
    Generate Word document version of the user guide
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title page
        title = doc.add_heading('🐄 Cattle Disease Detection System', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Comprehensive User Guide & System Documentation', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('AI-Powered Disease Detection for Modern Cattle Management').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # User info
        doc.add_paragraph(f'Generated for: {request.user.username}')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading('📋 Table of Contents', level=1)
        toc_items = [
            '1. System Overview',
            '2. Key Features', 
            '3. Getting Started',
            '4. User Interaction Flow',
            '5. User Interface Guide',
            '6. Device Compatibility',
            '7. Security & Privacy',
            '8. Benefits for Users',
            '9. Troubleshooting',
            '10. Support & Contact'
        ]
        
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
        
        doc.add_page_break()
        
        # System Overview
        doc.add_heading('🎯 System Overview', level=1)
        doc.add_paragraph(
            'The Cattle Disease Detection System is an advanced AI-powered web application designed to help farmers, '
            'veterinarians, and livestock managers detect diseases in cattle through intelligent image analysis. '
            'The system leverages cutting-edge machine learning algorithms to analyze uploaded images and identify '
            'potential foot and mouth infections in cattle with high accuracy.'
        )
        
        # Technology Foundation
        doc.add_heading('🔬 Technology Foundation', level=2)
        tech_table = doc.add_table(rows=1, cols=2)
        tech_table.style = 'Table Grid'
        hdr_cells = tech_table.rows[0].cells
        hdr_cells[0].text = 'Component'
        hdr_cells[1].text = 'Description'
        
        tech_data = [
            ('AI Engine', 'ONNX-based machine learning models'),
            ('Image Processing', 'Computer vision algorithms'),
            ('Web Platform', 'Django-based responsive web application'),
            ('Database', 'PostgreSQL for secure data storage'),
            ('Visualization', 'Multiple AI explanation techniques')
        ]
        
        for component, description in tech_data:
            row_cells = tech_table.add_row().cells
            row_cells[0].text = component
            row_cells[1].text = description
        
        doc.add_page_break()
        
        # Key Features
        doc.add_heading('⭐ Key Features', level=1)
        
        features = [
            ('🤖 AI-Powered Detection', 'Advanced computer vision models analyze cattle images with professional-grade accuracy'),
            ('⚡ Real-time Analysis', 'Get instant results with detailed explanations and visual indicators'),
            ('📊 Comprehensive Analytics', 'Dashboard with health statistics, trends, and interactive charts'),
            ('🔍 Visual Explanations', 'AI-generated heatmaps and analysis showing areas of concern'),
            ('📱 Mobile Responsive', 'Access from any device - desktop, tablet, or smartphone'),
            ('🔒 Secure & Private', 'Enterprise-grade security with encrypted data storage')
        ]
        
        for title, description in features:
            p = doc.add_paragraph()
            p.add_run(title).bold = True
            p.add_run(f'\n{description}')
        
        doc.add_page_break()
        
        # Getting Started
        doc.add_heading('🚀 Getting Started', level=1)
        
        doc.add_heading('📝 Account Registration', level=2)
        registration_steps = [
            'Navigate to the system URL and click "Create an account" on the login page',
            'Enter a unique username and create a strong password (minimum 8 characters)',
            'Confirm your password and review password strength indicators',
            'Your account is immediately active after successful registration'
        ]
        
        for step in registration_steps:
            doc.add_paragraph(step, style='List Number')
        
        doc.add_heading('🔐 System Login', level=2)
        doc.add_paragraph(
            'Access the system using your registered credentials. The login page features password visibility toggle '
            'and secure session management. Sessions automatically expire for security purposes.'
        )
        
        doc.add_page_break()
        
        # User Workflow
        doc.add_heading('👤 User Interaction Flow', level=1)
        
        doc.add_heading('🏠 Dashboard Overview', level=2)
        doc.add_paragraph(
            'Upon login, users are greeted with a comprehensive dashboard that provides real-time health metrics, '
            'interactive charts, quick actions, and recent activity updates.'
        )
        
        doc.add_heading('🐄 Cattle Registration Process', level=2)
        cattle_steps = [
            'Access "Register Cattle" from the sidebar menu',
            'Enter cattle information: Tag Number, Breed, and Age',
            'System validates unique tag numbers and required fields',
            'Successful registration displays confirmation and redirects to dashboard'
        ]
        
        for step in cattle_steps:
            doc.add_paragraph(step, style='List Number')
        
        doc.add_heading('🔬 Disease Detection Process', level=2)
        detection_steps = [
            'Upload image using drag & drop or file browser',
            'Select cattle from registered animals dropdown',
            'AI processes image and generates analysis',
            'View comprehensive results with visual explanations',
            'Take appropriate action based on findings'
        ]
        
        for step in detection_steps:
            doc.add_paragraph(step, style='List Number')
        
        doc.add_page_break()
        
        # Security & Privacy
        doc.add_heading('🔒 Security & Privacy', level=1)
        
        doc.add_heading('🛡️ Data Protection Measures', level=2)
        security_features = [
            'Encrypted Authentication: Secure login with password hashing',
            'Data Isolation: User data completely separated and protected',
            'Secure Image Storage: Uploaded images encrypted at rest',
            'HTTPS Encryption: All data transmission encrypted',
            'Session Management: Automatic logout for inactive sessions'
        ]
        
        for feature in security_features:
            doc.add_paragraph(feature, style='List Bullet')
        
        doc.add_page_break()
        
        # Benefits
        doc.add_heading('🎯 Benefits for Users', level=1)
        
        doc.add_heading('💰 Economic Benefits', level=2)
        economic_benefits = [
            'Cost Savings: Reduce veterinary costs through early detection',
            'Time Efficiency: Quick analysis compared to manual methods',
            'Productivity Gains: Maintain healthier herds',
            'Targeted Treatment: Focus resources where needed'
        ]
        
        for benefit in economic_benefits:
            doc.add_paragraph(benefit, style='List Bullet')
        
        doc.add_heading('🔬 Technical Benefits', level=2)
        technical_benefits = [
            'Early Detection: Identify diseases before visible symptoms',
            'Accuracy: AI analysis reduces human error',
            'Consistency: Standardized analysis across all cattle',
            'Documentation: Automatic digital health records'
        ]
        
        for benefit in technical_benefits:
            doc.add_paragraph(benefit, style='List Bullet')
        
        doc.add_page_break()
        
        # Support
        doc.add_heading('📞 Support & Contact', level=1)
        doc.add_paragraph(
            'For technical support, general inquiries, or training requests, please contact our support team. '
            'The system includes comprehensive help resources, video tutorials, and best practices guides.'
        )
        
        doc.add_heading('⚠️ Important Disclaimer', level=2)
        doc.add_paragraph(
            'This system is designed to assist in cattle health monitoring but should not replace professional '
            'veterinary diagnosis. Always consult with a qualified veterinarian for serious health concerns.'
        )
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph('Cattle Disease Detection System - Version 2.0')
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph('© 2024 Cattle Health Technologies. All rights reserved.')
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="Cattle_Disease_Detection_User_Guide_{datetime.now().strftime("%Y%m%d")}.docx"'
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        messages.error(request, "Error generating Word document. Please try again.")
        return redirect('dashboard')

@login_required
def help_center(request):
    """
    Help center page with download options
    """
    return render(request, 'detection/help_center.html')
